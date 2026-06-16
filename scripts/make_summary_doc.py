"""Generate a single-page, human-readable summary Word document for GRAPES-SHAP."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent.parent / "outputs" / "GRAPES_SHAP_Summary.docx"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

def heading(text):
    h = doc.add_paragraph()
    r = h.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
    h.space_after = Pt(2)
    return h

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("GRAPES-SHAP — Project Summary")
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("An interpretable, world-model-guided retrieval system for clinical question answering")
sr.italic = True
sr.font.size = Pt(10)

# Overview
heading("What this project is")
doc.add_paragraph(
    "GRAPES-SHAP is a medical question-answering system that goes beyond ordinary retrieval. "
    "It pulls the most relevant medical evidence, reasons about possible treatment plans using a small "
    "learned world model, estimates how confident it should be, and then explains which evidence mattered "
    "most using SHAP attributions. The final answer is written by a large language model that is conditioned "
    "on all of this structured reasoning."
)

# Datasets
heading("Datasets used")
doc.add_paragraph(
    "The system was built and tested on three well-known medical datasets, totalling about 100,000 samples:"
)
for t in [
    "DDXPlus — 80,000 training, 10,000 validation, and 10,000 test cases (diagnosis reasoning).",
    "MedMCQA — 50,000 documents used as the medical evidence corpus.",
    "MedQA — 1,000 test queries used for evaluation.",
]:
    doc.add_paragraph(t, style="List Bullet")

# Models
heading("Models used")
doc.add_paragraph(
    "Several models work together in a 12-step pipeline. Each one has a clear job:"
)
for t in [
    "Answer generator: DeepSeek 'deepseek-chat' large language model, fine-tuned with QLoRA.",
    "Evidence retrieval: a dense MiniLM + FAISS search combined with BM25 keyword search, merged by reciprocal-rank fusion.",
    "Re-ranking: a cross-encoder (ms-marco-MiniLM-L-6-v2) plus MMR to balance relevance and diversity.",
    "Query expansion: HyDE generates 3 helper sub-queries to improve recall.",
    "World model: a 3-layer GRU with a causal-residual block that simulates treatment outcomes.",
    "Knowledge graph: a causal graph with an edge-biased GNN for structured reasoning.",
    "Uncertainty: a 5-member deep ensemble that produces calibrated confidence.",
    "Planner: a Tree-of-Thought search over treatment actions.",
    "Interpretability: SHAP attributions show which evidence drove the answer.",
    "Safety: a Self-RAG and NLI hallucination check before the answer is accepted.",
]:
    doc.add_paragraph(t, style="List Bullet")

# Parameters
heading("Key parameters")
doc.add_paragraph(
    "Architecture: latent size 256, hidden size 512, observation dim 64, 50 actions, 20 graph nodes, "
    "5 outcomes, sequence length 8, 8 attention heads, 3 transformer layers, dropout 0.10, ensemble of 5."
)
doc.add_paragraph(
    "Retrieval & planning: top-k 6, embedding dim 384, SHAP permutations 32, MMR lambda 0.6, "
    "RRF k 60, HyDE sub-queries 3, planning horizon 4, 8 candidate plans."
)
doc.add_paragraph(
    "Language model: temperature 0.3, max tokens 2000, top-p 0.9, QLoRA rank 16 / alpha 32 / dropout 0.05."
)
doc.add_paragraph(
    "Training: world-model 15 epochs (lr 2e-4), predictor 10 epochs (lr 1e-3), batch size 64, "
    "gradient clip 1.0, weight decay 1e-4, mixed-precision FP16, seed 42, on an RTX 4080 SUPER GPU."
)
doc.add_paragraph(
    "The full model has only 10,130,060 trainable parameters and trained in about 18.8 minutes."
)

# Results
heading("Key results")
doc.add_paragraph("How well the world model predicts outcomes:")
for t in [
    "MAE 0.040 and RMSE 0.075 — very low prediction error.",
    "Accuracy 0.755 with calibration error (ECE) of just 0.030 and 1-sigma coverage of 0.80 — the confidence is honest, not overconfident.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph("How GRAPES-SHAP compares with a strong baseline RAG system over 10 complex clinical cases:")
for t in [
    "Clinical concept coverage rose from 0.70 to 0.97.",
    "Answer completeness rose from 0.50 to 0.84.",
    "Evidence citations more than doubled, from 2.0 to 5.1 on average.",
    "It adds SHAP evidence attribution (1.28), calibrated uncertainty, and treatment planning — none of which the baseline has.",
    "Its stated confidence (0.70) is lower but better calibrated than the baseline's 0.91.",
]:
    doc.add_paragraph(t, style="List Bullet")

# Takeaway
heading("In short")
doc.add_paragraph(
    "GRAPES-SHAP gives more complete, better-grounded, and more trustworthy medical answers than a "
    "conventional retrieval system, while staying small (about 10 million parameters) and fast to train "
    "(under 19 minutes). It also explains its reasoning and knows when to be uncertain."
)

doc.save(OUT)
print(f"Saved: {OUT}")
